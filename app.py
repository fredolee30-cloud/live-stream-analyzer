from flask import Flask, request, jsonify, send_file, g
from flask_cors import CORS
import os
import re
import docx
import PyPDF2
from io import BytesIO
import json
import uuid
from datetime import datetime, timedelta
import hashlib

app = Flask(__name__)
CORS(app)  # 允许跨域请求

# 配置
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
STORED_FOLDER = os.path.join(BASE_DIR, 'stored')
SHARES_FOLDER = os.path.join(BASE_DIR, 'shares')
CONFIG_FILE = os.path.join(BASE_DIR, 'config.json')

# Render 环境检测 - Render 免费版磁盘是临时的，使用内存存储
IS_RENDER_ENV = os.environ.get('RENDER_ENV', 'false') == 'true'

# 内存存储（用于 Render 环境）
render_file_storage = {}
render_share_storage = {}
render_config = {}

# 最大文件大小 50MB
MAX_FILE_SIZE = 50 * 1024 * 1024
# 最大存储文件数
MAX_STORED_FILES = 10

# 确保目录存在（非 Render 环境）
if not IS_RENDER_ENV:
    for folder in [UPLOAD_FOLDER, STORED_FOLDER, SHARES_FOLDER]:
        os.makedirs(folder, exist_ok=True)

# 加载配置
def load_config():
    default_config = {
        'api_keys': {
            'kimi': '',
            'deepseek': '',
            'qwen': ''
        },
        'default_model': 'moonshot-v1-32k'
    }
    # Render 环境使用内存配置
    if IS_RENDER_ENV:
        if not render_config:
            render_config.update(default_config)
        return render_config.copy()

    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            config = json.load(f)
            # 合并默认配置
            for key in default_config:
                if key not in config:
                    config[key] = default_config[key]
            return config
    return default_config

# 保存配置
def save_config(config):
    # Render 环境保存到内存
    if IS_RENDER_ENV:
        render_config.clear()
        render_config.update(config)
        return

    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


@app.route('/api/parse-file', methods=['POST'])
def parse_file():
    """解析上传的文件，提取文本内容"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400

        filename = file.filename.lower()
        file_content = file.read()

        if len(file_content) > MAX_FILE_SIZE:
            return jsonify({'error': f'文件大小超过限制 ({MAX_FILE_SIZE // 1024 // 1024}MB)'}), 400

        text = ''

        # 处理 TXT 文件
        if filename.endswith('.txt'):
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = file_content.decode('gbk')
                except:
                    text = file_content.decode('utf-8', errors='ignore')

        # 处理 Word 文件 (.docx)
        elif filename.endswith('.docx'):
            doc = docx.Document(BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n'.join(paragraphs)

        # 处理 PDF 文件
        elif filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            pages = []
            for page in pdf_reader.pages:
                pages.append(page.extract_text())
            text = '\n'.join(pages)

        # 处理其他文本格式
        elif filename.endswith('.md'):
            text = file_content.decode('utf-8')
        elif filename.endswith('.json'):
            # 尝试从 JSON 中提取文本（飞书/通义听悟导出格式）
            try:
                data = json.loads(file_content.decode('utf-8'))
                # 飞书妙记格式
                if 'transcript' in data:
                    text = data['transcript']
                # 通用 JSON 格式，尝试提取所有字符串值
                else:
                    text = json.dumps(data, ensure_ascii=False, indent=2)
            except:
                text = file_content.decode('utf-8')
        else:
            return jsonify({'error': '不支持的文件格式，请上传 .txt/.docx/.pdf/.md/.json 文件'}), 400

        # 清理文本
        text = clean_text(text)

        return jsonify({
            'success': True,
            'filename': filename,
            'text': text,
            'length': len(text),
            'message': f'成功解析 {len(text)} 字'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def clean_text(text):
    """清理文本，移除多余空白和乱码"""
    # 移除过多的空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 移除行首尾多余空格
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    # 移除常见的乱码字符
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    return text


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传并存储文件，最多存储 10 个"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400

        # Render 环境使用内存存储
        if IS_RENDER_ENV:
            if len(render_file_storage) >= MAX_STORED_FILES:
                return jsonify({
                    'error': f'已达到最大存储数量 ({MAX_STORED_FILES}个)',
                    'files': list(render_file_storage.keys()),
                    'action': 'need_delete'
                }), 400

            file_id = str(uuid.uuid4())
            file_content = file.read()

            # 解析文件内容
            text = parse_file_content(file.filename, file_content)
            if text is None:
                return jsonify({'error': '不支持的文件格式'}), 400

            text = clean_text(text)

            # 存储到内存
            render_file_storage[file_id] = {
                'filename': file.filename,
                'content': file_content,
                'text': text,
                'created_at': datetime.now().isoformat()
            }

            return jsonify({
                'success': True,
                'file_id': file_id,
                'filename': file.filename,
                'length': len(text),
                'text': text,
                'message': f'文件已存储，当前共 {len(render_file_storage)}/{MAX_STORED_FILES} 个文件'
            })

        # 非 Render 环境使用文件存储
        existing_files = os.listdir(STORED_FOLDER)
        if len(existing_files) >= MAX_STORED_FILES:
            return jsonify({
                'error': f'已达到最大存储数量 ({MAX_STORED_FILES}个)',
                'files': existing_files,
                'action': 'need_delete'
            }), 400

        file_id = str(uuid.uuid4())
        original_filename = file.filename
        stored_filename = f"{file_id}_{original_filename}"
        stored_path = os.path.join(STORED_FOLDER, stored_filename)

        file_content = file.read()
        with open(stored_path, 'wb') as f:
            f.write(file_content)

        text = parse_file_content(original_filename, file_content)
        if text is None:
            os.remove(stored_path)
            return jsonify({'error': '不支持的文件格式'}), 400

        text = clean_text(text)

        return jsonify({
            'success': True,
            'file_id': file_id,
            'filename': original_filename,
            'stored_filename': stored_filename,
            'length': len(text),
            'text': text,
            'message': f'文件已存储，当前共 {len(existing_files) + 1}/{MAX_STORED_FILES} 个文件'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def parse_file_content(filename, file_content):
    """解析文件内容，返回文本"""
    filename = filename.lower()
    text = ''

    if filename.endswith('.txt'):
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = file_content.decode('gbk')
            except:
                text = file_content.decode('utf-8', errors='ignore')
    elif filename.endswith('.docx'):
        doc = docx.Document(BytesIO(file_content))
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    elif filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
    elif filename.endswith('.md'):
        text = file_content.decode('utf-8')
    elif filename.endswith('.json'):
        try:
            data = json.loads(file_content.decode('utf-8'))
            text = data.get('transcript', json.dumps(data, ensure_ascii=False, indent=2))
        except:
            text = file_content.decode('utf-8')
    else:
        return None

    return text


@app.route('/api/files', methods=['GET'])
def list_files():
    """列出已存储的文件"""
    try:
        # Render 环境使用内存存储
        if IS_RENDER_ENV:
            files = [
                {
                    'file_id': file_id,
                    'filename': info['filename'],
                    'size': len(info['content']),
                    'created_at': info['created_at']
                }
                for file_id, info in render_file_storage.items()
            ]
            return jsonify({'files': files, 'count': len(files), 'max': MAX_STORED_FILES})

        # 非 Render 环境使用文件存储
        files = []
        for filename in os.listdir(STORED_FOLDER):
            filepath = os.path.join(STORED_FOLDER, filename)
            stat = os.stat(filepath)
            files.append({
                'file_id': filename.split('_', 1)[0] if '_' in filename else filename,
                'filename': '_'.join(filename.split('_', 1)[1:]) if '_' in filename else filename,
                'size': stat.st_size,
                'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat()
            })
        return jsonify({'files': files, 'count': len(files), 'max': MAX_STORED_FILES})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/files/<file_id>', methods=['DELETE'])
def delete_file(file_id):
    """删除指定文件"""
    try:
        # Render 环境使用内存存储
        if IS_RENDER_ENV:
            if file_id in render_file_storage:
                del render_file_storage[file_id]
                return jsonify({'success': True, 'message': '文件已删除'})
            return jsonify({'error': '文件不存在'}), 404

        # 非 Render 环境使用文件存储
        for filename in os.listdir(STORED_FOLDER):
            if filename.startswith(f"{file_id}_"):
                filepath = os.path.join(STORED_FOLDER, filename)
                os.remove(filepath)
                return jsonify({'success': True, 'message': '文件已删除'})
        return jsonify({'error': '文件不存在'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/files/clear', methods=['POST'])
def clear_files():
    """清空所有存储文件"""
    try:
        # Render 环境使用内存存储
        if IS_RENDER_ENV:
            render_file_storage.clear()
            return jsonify({'success': True, 'message': '已清空所有文件'})

        # 非 Render 环境使用文件存储
        for filename in os.listdir(STORED_FOLDER):
            filepath = os.path.join(STORED_FOLDER, filename)
            os.remove(filepath)
        return jsonify({'success': True, 'message': '已清空所有文件'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/config', methods=['GET'])
def get_config():
    """获取配置（不返回敏感信息）"""
    config = load_config()
    # 隐藏 API Key 的具体值
    api_keys_status = {}
    for key, value in config.get('api_keys', {}).items():
        api_keys_status[key] = 'configured' if value else 'not_configured'
    return jsonify({
        'api_keys_status': api_keys_status,
        'default_model': config.get('default_model', 'moonshot-v1-32k')
    })


@app.route('/api/config', methods=['POST'])
def update_config():
    """更新配置"""
    try:
        data = request.json
        config = load_config()

        if 'api_keys' in data:
            for key, value in data['api_keys'].items():
                if value:  # 只更新非空值
                    config['api_keys'][key] = value

        if 'default_model' in data:
            config['default_model'] = data['default_model']

        save_config(config)
        return jsonify({'success': True, 'message': '配置已更新'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/analyze', methods=['POST'])
def analyze():
    """调用 AI 模型分析文稿"""
    try:
        data = request.json
        script = data.get('script', '')
        model = data.get('model', 'moonshot-v1-32k')
        api_key = data.get('api_key', '')

        if not script:
            return jsonify({'error': '文稿内容为空'}), 400

        # 如果没有传入 API Key，使用配置的
        if not api_key:
            config = load_config()
            if model.startswith('moonshot'):
                api_key = config.get('api_keys', {}).get('kimi', '')
            elif model.startswith('deepseek'):
                api_key = config.get('api_keys', {}).get('deepseek', '')
            elif model.startswith('qwen'):
                api_key = config.get('api_keys', {}).get('qwen', '')

        if not api_key:
            return jsonify({'error': '请先配置 API Key'}), 400

        # 确定 API 端点
        if model.startswith('moonshot'):
            api_url = 'https://api.moonshot.cn/v1/chat/completions'
        elif model.startswith('deepseek'):
            api_url = 'https://api.deepseek.com/v1/chat/completions'
        elif model.startswith('qwen'):
            api_url = 'https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation'
        else:
            return jsonify({'error': '不支持的模型'}), 400

        # 构建请求
        import requests

        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}'
        }

        # 构建分析提示词
        prompt = build_analysis_prompt(script)

        if model.startswith('qwen'):
            request_body = {
                'model': model,
                'input': {
                    'messages': [{'role': 'user', 'content': prompt}]
                },
                'parameters': {
                    'temperature': 0.3,
                    'result_format': 'json'
                }
            }
        else:
            request_body = {
                'model': model,
                'messages': [{'role': 'user', 'content': prompt}],
                'temperature': 0.3,
                'response_format': {'type': 'json_object'}
            }

        response = requests.post(api_url, headers=headers, json=request_body, timeout=60)

        if not response.ok:
            error_text = response.text
            try:
                error_data = response.json()
                error_text = error_data.get('error', {}).get('message', error_text)
            except:
                pass
            return jsonify({'error': f'API 调用失败：{error_text}'}), response.status_code

        result = response.json()

        # 解析响应
        if model.startswith('qwen'):
            content = result.get('output', {}).get('text', '')
        else:
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')

        # 尝试解析 JSON
        try:
            json_str = content.strip()
            json_match = re.search(r'\{[\s\S]*\}', json_str)
            if json_match:
                json_str = json_match.group(0)
            analysis_result = json.loads(json_str)
            return jsonify({'success': True, 'result': analysis_result})
        except json.JSONDecodeError as e:
            return jsonify({'error': f'解析结果失败：{str(e)}', 'raw_content': content}), 500

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def build_analysis_prompt(script):
    """构建分析提示词"""
    return f"""你是一位专业的金融财经直播内容分析专家。请对以下直播文稿进行全面分析，严格按照 JSON 格式输出结果。

【分析维度】

1. 内容主题分析（6 大类，权重总和 100%）
   - 宏观经济：经济数据（GDP/CPI/PPI/PMI/就业/社融）、货币政策（利率/准备金率/央行操作）、财政政策、汇率外汇、全球经济
   - 科技资产：AI、半导体、云计算、新能源、电动车、生物科技、消费电子、互联网平台
   - 大宗商品：贵金属（黄金/白银）、能源（原油/天然气）、工业金属（铜/铝/铁矿石）、农产品
   - 大类资产配置：权益资产（股票/股基）、固定收益（债券/债基/理财）、另类投资、现金管理、保险规划
   - 金融投教知识：投资基础、分析方法、风险管理、投资心理、理财规划
   - 直播互动：闲聊、开场白、问答互动、过渡语等非实质内容

2. 合规性检测（6 类违规）
   - 承诺收益："稳赚不赔"、"一定涨"、"保本保息"等
   - 预测涨跌："明天必涨"、"目标价 XX"、"抄底时机"等
   - 政治议题：敏感政治人物、政策猜测、地缘政治判断等
   - 推荐产品："买 XX 股票"、"推荐 XX 基金"、"重仓 XX"等
   - 诱导投资：催促下单、限时优惠、制造焦虑等
   - 违反广告法："最"、"第一"、"绝对"、"100%"等绝对化用语

3. 内容质量评估（9 维度，各 0-100 分，按权重计算总分）
   - 需求贴合度（权重 18%）：话题是否贴合市场需求、从用户痛点出发、解决核心问题
   - 论述质量（权重 18%）：问题界定是否清晰、观点是否明确、论证过程是否有说服力
   - 说服力（权重 15%）：论证深度、反面观点处理
   - 逻辑完整性（权重 12%）：论点 - 论据 - 结论链条是否完整
   - 实用价值（权重 12%）：对听众投资理财是否有实际指导意义和可操作性
   - 逻辑框架（权重 10%）：整体结构清晰度、话题引入/总结、过渡衔接
   - 易懂性（权重 7%）：专业概念是否"降维"表达、复杂原理是否通俗易懂
   - 数据详实度（权重 3%）：具体数据引用数量和质量
   - 案例质量（权重 5%）：案例相关性、时效性、典型性

【输出格式】
请严格输出以下 JSON 格式（不要有任何额外说明）：

{{
    "theme": {{
        "macro_weight": 0-100,
        "tech_weight": 0-100,
        "commodity_weight": 0-100,
        "allocation_weight": 0-100,
        "education_weight": 0-100,
        "interaction_weight": 0-100,
        "summary": "主题分析总结"
    }},
    "compliance": {{
        "is_compliant": true/false,
        "issues": [
            {{
                "type": "承诺收益/预测涨跌/政治议题/推荐产品/诱导投资/违反广告法",
                "content": "原文引用",
                "severity": "high/medium/low",
                "suggestion": "修改建议"
            }}
        ],
        "summary": "合规性总结"
    }},
    "quality": {{
        "need_fit_score": 0-100,
        "argument_score": 0-100,
        "persuasion_score": 0-100,
        "logic_score": 0-100,
        "practical_score": 0-100,
        "framework_score": 0-100,
        "clarity_score": 0-100,
        "data_score": 0-100,
        "case_score": 0-100,
        "overall_score": 0-100,
        "deductions": [
            {{
                "dimension": "维度名称",
                "issue": "问题描述",
                "suggestion": "改进建议"
            }}
        ]
    }},
    "suggestions": [
        "具体改进建议 1",
        "具体改进建议 2"
    ]
}}

【直播文稿内容】
{script}"""


@app.route('/api/share/create', methods=['POST'])
def create_share():
    """创建分享链接"""
    try:
        data = request.json
        analysis_result = data.get('result', {})
        script = data.get('script', '')

        if not analysis_result:
            return jsonify({'error': '分析结果为空'}), 400

        # 生成分享 ID
        share_id = str(uuid.uuid4())[:8]
        share_data = {
            'share_id': share_id,
            'created_at': datetime.now().isoformat(),
            'expires_at': (datetime.now() + timedelta(days=7)).isoformat(),
            'script': script,
            'result': analysis_result
        }

        # Render 环境使用内存存储
        if IS_RENDER_ENV:
            render_share_storage[share_id] = share_data
            return jsonify({
                'success': True,
                'share_id': share_id,
                'share_url': f"/share/{share_id}",
                'expires_at': share_data['expires_at']
            })

        # 非 Render 环境使用文件存储
        share_file = os.path.join(SHARES_FOLDER, f"{share_id}.json")
        with open(share_file, 'w', encoding='utf-8') as f:
            json.dump(share_data, f, ensure_ascii=False, indent=2)

        return jsonify({
            'success': True,
            'share_id': share_id,
            'share_url': f"/share/{share_id}",
            'expires_at': share_data['expires_at']
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/share/<share_id>', methods=['GET'])
def get_share(share_id):
    """获取分享内容"""
    try:
        # Render 环境使用内存存储
        if IS_RENDER_ENV:
            if share_id not in render_share_storage:
                return jsonify({'error': '分享链接不存在'}), 404

            share_data = render_share_storage[share_id]

            # 检查是否过期
            expires_at = datetime.fromisoformat(share_data['expires_at'])
            if datetime.now() > expires_at:
                del render_share_storage[share_id]
                return jsonify({'error': '分享链接已过期'}), 404

            return jsonify({
                'success': True,
                'data': {
                    'script': share_data['script'],
                    'result': share_data['result'],
                    'created_at': share_data['created_at'],
                    'expires_at': share_data['expires_at']
                }
            })

        # 非 Render 环境使用文件存储
        share_file = os.path.join(SHARES_FOLDER, f"{share_id}.json")

        if not os.path.exists(share_file):
            return jsonify({'error': '分享链接不存在'}), 404

        with open(share_file, 'r', encoding='utf-8') as f:
            share_data = json.load(f)

        # 检查是否过期
        expires_at = datetime.fromisoformat(share_data['expires_at'])
        if datetime.now() > expires_at:
            # 删除过期文件
            os.remove(share_file)
            return jsonify({'error': '分享链接已过期'}), 404

        return jsonify({
            'success': True,
            'data': {
                'script': share_data['script'],
                'result': share_data['result'],
                'created_at': share_data['created_at'],
                'expires_at': share_data['expires_at']
            }
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health():
    """健康检查"""
    return jsonify({'status': 'ok', 'message': '服务运行中'})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
